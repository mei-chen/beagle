# -*- coding: utf-8 -*-

"""
    cleanup
    ==========

    A cleanup-related routines
"""

import logging
import os
import re
import shutil
import tempfile
import zipfile
from collections import namedtuple
from difflib import SequenceMatcher
from itertools import islice
from string import ascii_lowercase, punctuation, ascii_uppercase

import docx
import inflect
from bs4 import BeautifulSoup
from docx.document import Document
from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.table import CT_Tbl, CT_TblGrid
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from lxml.etree import fromstring, tostring

logger = logging.getLogger(__name__)
inflect = inflect.engine()

RE_LIST = re.compile(
    u'^\s*([\(\[]?'
    u'([1-9][0-9]*|[mcxvil]+|'
    u'[a-z*:\u2014\u2022\u2023\u25E6\u2043\u2219\uf0b7-])'
    u'\s*([\)\]\.\:]|(\.[0-9a-z]*))+|'
    u'[*:\u2014\u2022\u2023\u25E6\u2043\u2219\uf0b7-]{1,5})(\s+|$)',
    re.I | re.U
)
RE_LIST_2 = re.compile(
    u'^\s*[\(\[]?([a-z])\\1{1,2}[\)\]\.\:](\s+|$)',
    re.I | re.U
)
RE_NOT_LIST = re.compile(
    u'^\s*N.B.',
    re.I | re.U
)

RE_PERIOD = re.compile(r'\.$')
RE_SPACE = re.compile(r'^\s*', re.U)


def is_not_newline(line, prev_line):
    stripped_prev_line = ListCleanup.strip_list_prefix(prev_line).strip()
    # If prev line is just list item - wrap it with next line
    if len(prev_line) and not len(stripped_prev_line):
        return True
    # New line - only if prev line ends with .?! and first char is uppercase
    return bool(
        len(stripped_prev_line) and len(line) and
        stripped_prev_line[-1] not in '.?!' and
        not line[0].isupper()
    )


class ListCleanup(object):
    def __init__(self, doc, text_lines):
        self.doc = doc
        self.last_line = -1
        self.text_lines = []
        self.reset_list_data()
        self.text_lines = text_lines

    @classmethod
    def is_list_elem(cls, elem):
        try:
            return bool(
                ((RE_LIST.match(elem.text) or RE_LIST_2.match(elem.text)) and
                 not RE_NOT_LIST.match(elem.text)) or
                (elem._p.pPr and elem._p.pPr.numPr)
            )
        except AttributeError:
            return None

    @classmethod
    def strip_list_prefix(cls, text):
        # If we didn't find first pattern - use second pattern
        text1 = RE_LIST.sub('', text)
        if text1 == text:
            return RE_LIST_2.sub('', text)
        return text1

    def reset_list_data(self):
        self.level_margin = None
        self.prev_paragraph = None
        self.prepending_paragraph = None
        self.prepending_text = ''
        self.prev_text = ''

    def is_top_level_element(self, paragraph):
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            # Check if formatted list item is top-level
            return paragraph._p.pPr.numPr.ilvl.val == 0
        margin = RE_SPACE.findall(paragraph.text)[0]
        margin.replace('\t', '        ')
        level_margin = len(margin)
        if not self.level_margin:
            self.level_margin = level_margin
        return level_margin <= self.level_margin

    def get_paragraph_prefix(self, paragraph):
        # If this isn't a formatted list - do nothing
        if paragraph._p.pPr is None or paragraph._p.pPr.numPr is None:
            return None
        text = paragraph.text.strip(' \t')
        # search for text among array
        for i in range(self.last_line, len(self.text_lines)):
            try:
                pos = self.text_lines[i].index(text)
                self.last_line = i
                return self.text_lines[i][:pos].strip(' \t')
            # no substring found in line - probe next
            except ValueError:
                pass
        # nothing is found
        return None

    def clear_paragraph_list(self, paragraph):
        if paragraph._p.pPr is None or paragraph._p.pPr.numPr is None:
            return
        paragraph._p.pPr.numPr.clear()

    def process_list_item_cleanup(self, paragraph, text):
        # get formatted prefix if plaintext
        prefix = self.get_paragraph_prefix(paragraph)
        # append formatted prefix
        if prefix:
            text = ' '.join((prefix, text))
        if self.prev_paragraph:
            self.prev_text = self.prev_paragraph.text.rstrip(' \t')
            # append text to prev paragraph if exists
            self.prev_text = ' '.join((self.prev_text, text))
            self.prev_paragraph.text = self.prev_text
            del_paragraph(paragraph)
        else:
            # or save to current paragraph
            paragraph.text = text
            # and clear list attr
            self.clear_paragraph_list(paragraph)
            self.prev_paragraph = paragraph

    def process_list_item_rewrite(self, paragraph, text):
        text = self.strip_list_prefix(text)
        if self.prepending_text:
            # append text to prev non-list paragraph if exists
            text = ' '.join((self.prepending_text, text))
        if self.prepending_paragraph:
            del_paragraph(self.prepending_paragraph)
            self.prepending_paragraph = None
        # save text to current paragraph
        paragraph.text = text
        # and clear list attr
        self.clear_paragraph_list(paragraph)
        self.prev_paragraph = paragraph

    def list_cleanup(self, rewriting=False):
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip(' \t')
            # remove empty line
            if not text:
                del_paragraph(paragraph)
                continue
            # skip non-list pars
            if not self.is_list_elem(paragraph):
                self.reset_list_data()
                # do NOT attach to paragraps that ends with .!?
                if text[-1] in ['.', '?', '!']:
                    continue
                self.prev_paragraph = paragraph
                self.prepending_paragraph = paragraph
                self.prepending_text = paragraph.text.rstrip(' \t')
                continue
            if rewriting and self.is_top_level_element(paragraph):
                self.process_list_item_rewrite(paragraph, text)
            else:
                self.process_list_item_cleanup(paragraph, text)
        return self.doc


class TOCCleanupException(Exception):
    pass


class TOCCleanup(object):
    # <orderNum> <Text> <pageNum>
    TOC_ITEM_1 = re.compile(u'^\s*(\d+)[\d\.]{0,4}.*(?:\s|\.{4,}(?:\s?\.))\d+$')
    # <orderNum> <Text>
    TOC_ITEM_2 = re.compile(u'^\s*(\d+)[\.\d]{0,7}.*[^\d]$')
    # <Text> <pageNum>
    TOC_ITEM_3 = re.compile(u'^\s*[^\d][A-Za-z]+[^.]{5,}\s+(\d{1,3})(?:\.\d+){0,3}$')
    # <Text> ..... <pageNum>
    TOC_ITEM_4 = re.compile(u'^[^\d].{5,}\s?\.{4,}\s*(\d{1,3})$')

    # TOC element can contain one of those styles
    TOC_STYLE_PARTIAL_NAMES = ['toc']
    TOC_NAMES = ['table of contents', 'toc', 'contents', 'index']

    ITEM_CAN_CONTAIN = {'appendix', 'annex'}
    _cached_items = None
    # First valid TOC item will be stored here as namedtuple
    # during cached_items call or during get_toc_start call
    _toc_start = None
    # Indicates is cached_items has all items from gen_precleaned_items
    # if during cached_items call the _preclean_toc_validator method returns True
    # then see comments above and this one will remain unchanged
    _is_fully_cached = False
    # During _bulk_check method call this one goes to False. Should give more stability but seems like
    # it does not much
    _check_page_num = True
    # Other flags, nums for the plain text toc cleanup
    __last_page_num = 0
    __last_page_num_did_fix = False
    __last_item = None

    def __init__(self, name, doc):
        self.name = name
        self.doc = doc
        self.plain_validator = self.plain_primary_validator
        self.errors = self.Errors()

    class Errors:
        rate = 0.0
        threshold = 10.0

        def calculate(self, item_nt):
            text = item_nt.item.text
            text_no_dots = text.translate({ord('.'): None})
            if len(text_no_dots) > 50:
                return (2 + self.rate) / self.threshold
            return 0.1

        @property
        def threshold_reached(self):
            return self.rate >= 1.0

        def reset(self):
            self.rate = 0.0

    # PREPROCESS
    @property
    def cached_items(self):
        """
        Makes list of docx paragraphs
        If during it method will get a TOC keyword it will break iterations and
        set TOC start index to __cached_items_toc_start
        Returns: docx paragraphs
        """
        if self._cached_items is None:
            items = []

            for item_nt in self.gen_precleaned_items():
                items.append(item_nt)
                if item_nt.is_table:
                    continue

                if self._preclean_toc_validator(item_nt) is not None:
                    self._toc_start = item_nt
                    self._cached_items = items
                    self._is_fully_cached = False
                    return self._cached_items
            self._cached_items = items
            self._is_fully_cached = True

        return self._cached_items

    def gen_precleaned_items(self):
        doc = self.doc
        # index in iter_block_items(docx.Document) (unlike docx.paragraphs there tables are included)
        real_index = -1
        # index in self.cached_items
        index = -1
        nt = namedtuple('Item', ['index', 'real_index', 'item', 'is_table'])
        parent_elm = self.__doc_get_parent_element()

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                real_index += 1
                paragraph = Paragraph(child, doc)
                if self._preclean_item_validator(paragraph):
                    index += 1
                    yield nt(index, real_index, Paragraph(child, doc), False)
            elif isinstance(child, CT_Tbl):
                real_index += 1
                index += 1
                yield nt(index, real_index, Table(child, doc), True)

    def __doc_get_parent_element(self):
        doc = self.doc
        if isinstance(doc, Document):
            return doc.element.body
        elif isinstance(doc, _Cell):
            return doc._tc
        raise ValueError('Cannot get items from the document "%s"' % self.name)

    def _preclean_item_validator(self, paragraph):
        if self._is_generated_toc_item(paragraph):
            return True
        text = paragraph.text.strip()

        preconj = (
            'the', 'as', 'of', 'its', 'a', 'be', 'and', 'by', 'at', 'or', 'are',
            '!', '#', '$', '%', '&', '*', '+', ',', '-', '.', '/', ':', ';', '<', '=', '>', '?', '@',
            '\\', '^', '_', '`', '}'
        )
        validators = (
            lambda: not text.startswith(ascii_lowercase),
            lambda: not text.endswith(preconj),
        )

        return 20 > len(text.split()) >= 1 and all(validator() for validator in validators)

    def _preclean_toc_validator(self, item_nt):
        text = item_nt.item.text
        clean_text = text.translate({ord(ch): None for ch in '0123456789.'})
        lower_text = clean_text.strip().lower()

        if len(lower_text) < 30 and any(toc_name in lower_text for toc_name in TOCCleanup.TOC_NAMES):
            return True
        elif (self.TOC_ITEM_1.match(text) or self.TOC_ITEM_4.match(text)) and self.is_mixed_case(clean_text):
            return True

    def get_items(self, start=0, stop=None, include_tables=False):
        """
        Args:
            start: optional index to start from
            stop: optional stop index
            include_tables: filters out the Table instances

        Returns: cached items or generator (depends on the self.is_fully_cached flag)

        """
        if self._is_fully_cached:
            if include_tables:
                return self.cached_items[start:stop]
            else:
                return (item for item in self.cached_items[start:stop] if not item.is_table)
        else:
            if include_tables:
                return islice(self.gen_precleaned_items(), start, stop)
            else:
                return (item for item in islice(self.gen_precleaned_items(), start, stop) if not item.is_table)

    # GENERAL VALIDATORS
    def _match_any_pattern(self, item_nt):
        patterns = (
            self.TOC_ITEM_1.match,
            self.TOC_ITEM_2.match,
            self.TOC_ITEM_3.match,
            self.TOC_ITEM_4.match
        )
        return any(pattern(item_nt.item.text) for pattern in patterns)

    def _is_generated_toc_item(self, item):
        # fairy of imagination says no
        # This method will be called or during cached_items invocation then the argument "item" could be an instance
        # of docx.Paragraph, or during cleanup and the argument "item" will be an instance of the namedtuple.
        item = getattr(item, 'item', item)
        if item.text == 'to update table of contents - press F9':
            return True
        return any(style_name in item.style.name.lower() for style_name in self.TOC_STYLE_PARTIAL_NAMES)

    def _is_listing_item(self, item_nt):
        return item_nt.item._p.pPr and bool(item_nt.item._p.pPr.numPr)

    @staticmethod
    def is_mixed_case(text):
        uppers = 0
        lowers = 0
        words = text.split(' ')
        for word in words:
            if word.istitle() or word.isupper():
                uppers += 1
            else:
                lowers += 1
        return uppers > lowers

    # TOC START SEARCH METHODS
    def get_toc_start(self):
        """
        self.__toc_start
        This attribute could be set during self.cached_items call
        in this case we found a TOC for sure. But be careful, in this case the cached_items property has only items
        before TOC start index

        _check_items_duplications can return None, True. In this case found all N duplications
        or it can return (index, item), False. In this case during duplication search the algorithm meet a string
        looks like a TOC item
        """
        items_to_check = []
        for item_nt in self.cached_items:
            if self._toc_start:
                return
            elif len(items_to_check) < 5:
                items_to_check.append(item_nt)
                continue
            self._check_page_num = False
            self._toc_start = self._bulk_check(items_to_check)
            self._check_page_num = True
            items_to_check = [item_nt]
        raise TOCCleanupException('Cannot get TOC Start')

    def _bulk_check(self, items):
        duplicates = self.__bulk_check_duplications(items)
        for d1, d2 in duplicates:
            d1_valid_forward = self.plain_seek_forward(d1.index, d1.index + 10, ignore_threshold=True)
            if d1_valid_forward:
                return d1
            d2_valid_forward = self.plain_seek_forward(d2.index, d1.index + 10, ignore_threshold=True)
            if d2_valid_forward:
                return d2

    @staticmethod
    def match_sequences(seq1, seq2, junk=punctuation + '0123456789'):
        if junk:
            table = {ord(ch): None for ch in junk}
            seq1 = seq1.translate(table)
            seq2 = seq2.translate(table)
        seq1 = seq1.lower().strip()
        seq2 = seq2.lower().strip()
        return SequenceMatcher(None, seq1, seq2)

    def __bulk_check_duplications(self, left_items):
        """
        Takes items to compare (left_items) and match each one with current item (right_item)
        Returns: list of duplicates (if any) [(left_item, right_item)]
        """
        duplicates = []
        duplicate_indexes = set()
        for right_item in self.get_items():
            for left_item in left_items:
                left_text = left_item.item.text
                right_text = right_item.item.text
                is_valid_pair = all((
                    right_item.real_index != left_item.real_index,
                    self.match_sequences(right_text, left_text).ratio() >= 0.95
                ))
                if is_valid_pair and not duplicate_indexes.intersection({right_item.index, left_item.index}):
                    duplicates.append((left_item, right_item))
                    duplicate_indexes.update({right_item.index, left_item.index})
                    left_items.remove(left_item)
                elif len(duplicates) >= 2:
                    return duplicates
        return duplicates

    def get_toc_type(self):
        """
        Depends on toc type runs an appropriate cleanup
        Returns: (<namedtuple>, <cleanup_method>)
        """
        start = self._toc_start.index
        for item_nt in self.get_items(start=start, stop=start + 10, include_tables=True):
            if item_nt.is_table:
                return item_nt, self.run_table
            elif self._is_generated_toc_item(item_nt):
                return item_nt, self.run_generated
            elif self._match_any_pattern(item_nt):
                return item_nt, self.run_plain
            elif self._is_listing_item(item_nt):
                return item_nt, self.run_listing
        raise TOCCleanupException('Cannot get TOC Type.')

    # CLEANUP ENTRY POINTS
    def run_table(self, **kwargs):
        """
        Entry point TOC Table cleanup
        Just remove TOC Title and Table
        """
        del_paragraph(self._toc_start.item)
        del_paragraph(kwargs['first_item'].item)

    def run_generated(self, **kwargs):
        """
        Generated TOC items accessibly only thorough xml otherwise (like access via
        doc.paragraphs[<num>]) empty content will be returned
        """
        del_paragraph(self._toc_start.item)
        for item_nt in self.get_items(self._toc_start.index):
            if not self._is_generated_toc_item(item_nt.item):
                break
            del_paragraph(item_nt.item)

    def run_plain(self, **kwargs):
        start = self._toc_start.index
        end = self.plain_seek_forward(start)
        if start is None or end is None or start > end:
            raise TOCCleanupException('Invalid Indexes. Start index {} End index {}'.format(start, end))
        for item_nt in self.get_items(start, end + 1):
            del_paragraph(item_nt.item)

    def run_listing(self, **kwargs):
        del_paragraph(self._toc_start.item)
        for item_nt in self.get_items(self._toc_start.index):
            if not self._is_listing_item(item_nt):
                break
            del_paragraph(item_nt.item)
        last_valid_index = self.plain_seek_forward(self._toc_start.index)
        if last_valid_index:
            for item_nt in self.get_items(self._toc_start.index, last_valid_index + 1):
                del_paragraph(item_nt.item)

    # PLAIN TEXT RELATED VALIDATORS
    def fix_last_page_num(self, cur_page_num, item_nt):
        match = self.TOC_ITEM_1.match(item_nt.item.text)
        if bool(match):
            self.__last_page_num = cur_page_num
            return True
        elif not self.__last_page_num_did_fix and (item_nt.index - self._toc_start.index) < 3:
            self.__last_page_num_did_fix = True
            return True

    def plain_check_page_num(self, cur_page_num, item_nt):
        if self.errors.rate > 0.5 and self.fix_last_page_num(cur_page_num, item_nt):
            return True
        prev_page_num = self.__last_page_num
        gte = cur_page_num >= prev_page_num
        diff = cur_page_num - prev_page_num
        return diff < 20 and gte

    def plain_primary_validator(self, item_nt):
        text = item_nt.item.text
        validators = (
            self.TOC_ITEM_1.match,
            self.TOC_ITEM_2.match,
            self.TOC_ITEM_3.match,
            self.TOC_ITEM_4.match
        )
        for validator in validators:
            match = validator(text)
            if bool(match):
                if not self._check_page_num:
                    return True
                cur_page_num = int(match.group(1))
                is_gte = self.plain_check_page_num(cur_page_num, item_nt)
                if is_gte:
                    self.__last_page_num = cur_page_num
                    return True
        return self.plain_secondary_validator(item_nt)

    def plain_secondary_validator(self, item_nt):
        current_text = item_nt.item.text
        if any((len(current_text.strip()) > 90, RE_LIST.match(current_text))):
            return False
        current_words = current_text.strip().lower().split(' ')
        if any(w in self.ITEM_CAN_CONTAIN for w in current_words):
            return True
        elif self.__last_item is not None:
            previous_words = self.__last_item.item.text.strip().lower().split(' ')
            if current_words[0] == previous_words[0]:
                return True
        return current_text.count('.') > 15

    # PLAIN TEXT RELATED HELPERS
    def plain_seek_forward(self, start, stop=None, ignore_threshold=False):
        last_valid_index = None
        for item_nt in self.get_items(start, stop):
            is_valid = self.plain_validator(item_nt)
            if not is_valid:
                if ignore_threshold:
                    return False
                elif self.errors.threshold_reached:
                    self.errors.reset()
                    return last_valid_index
                else:
                    self.errors.rate += self.errors.calculate(item_nt)
            else:
                last_valid_index = item_nt.index
                self.errors.reset()
            self.__last_item = item_nt
        return last_valid_index

    def run(self):
        self.get_toc_start()
        first_valid_item, cleanup_method = self.get_toc_type()
        cleanup_method(first_item=first_valid_item)
        return self.doc


def del_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None


def get_separator(sentence):
    """
    Checks if all words in sentence are singular nouns.
    Special case is sentences with `:` in them. In that case we return `that`
    Returns 'is', 'are' or `that.
    """

    if ':' in sentence:
        return 'that'

    for w in sentence.split():
        # returns string in case w is plural.
        # Otherwise (w is singular or isn't a noun) returns False
        if inflect.singular_noun(w):
            return 'are'

    return 'is'


def process_highlight(content, highlight_matrix):
    """
    Changes table content according to highlighted cells (choices)
    """

    def _uneven_type_1(content):
        """
        Check if table type is:
            .----------.-----------.-----------.
            |          | Choice1.1 | Choice1.2 |
            |          +-----------+-----------:
            | Question | Choice2.1 | Choice2.2 |
            |          +-----------+-----------:
            |          | Choice3.1 | Choice3.2 |
            '----------'-----------'-----------'
        """

        prev_row = content[0]
        for row in content:
            if row[0] != prev_row[0]:
                return False
            prev_row = row

        return True

    if not any([any(row) for row in highlight_matrix]):
        return content

    uneven_type_1 = _uneven_type_1(content)

    is_symmetric = all([len(row) == len(content[0]) for row in content])

    if is_symmetric and uneven_type_1:
        new_content = [[content[0][0]], ]
        for row_i in range(len(content)):
            for cell_i in range(1, len(content[row_i])):
                if highlight_matrix[row_i][cell_i]:
                    new_content[0].append(content[row_i][cell_i])

        return new_content

    ## Special cases
    # .-----------------------------------.
    # |             Question              |
    # :-----------+-----------+-----------:
    # | Choice1.1 | Choice1.2 | Choice1.3 |
    # :-----------+-----------+-----------:
    # | Choice2.1 | Choice2.2 | Choice2.3 |
    # '-----------'-----------'-----------'
    elif not is_symmetric and len(content[0]) == 1 and len(content) > 1:
        new_content = []
        header = content[0][0]
        for row_i in range(1, len(content)):
            if len(content[row_i]) == 1:
                # New title in the middle of the table
                header = content[row_i][0]
                continue

            elif not any(highlight_matrix[row_i]):
                new_content.append(content[row_i])
                continue

            new_row = [header]
            for cell_i in range(len(content[row_i])):
                if highlight_matrix[row_i][cell_i]:
                    new_row.append(content[row_i][cell_i])

            new_content.append(new_row)

        return new_content

    # Mixed format
    # :-----------+-----------+-----------:
    # |   text    |   text    |   text    |
    # :-----------+-----------+-----------:
    # | Question  | Choice 1  | Choice 2  |
    # .-----------------------------------.
    # |              Text                 |
    # :-----------+-----------+-----------:
    # | Question  | Choice 1  | Choice 2  |
    # :-----------+-----------+-----------:
    # |      text       |      text       |
    # '-----------'-----------'-----------'
    else:
        new_content = []
        for row_i in range(len(content)):
            if not any(highlight_matrix[row_i]):
                new_content.append(content[row_i])
                continue

            new_row = []
            for cell_i in range(len(content[row_i])):
                if cell_i == 0:
                    new_row.append(content[row_i][0])
                elif highlight_matrix[row_i][cell_i]:
                    new_row.append(content[row_i][cell_i])
            new_content.append(new_row)

        return new_content

    return content

class TablePostprocessing(object):
    def __init__(self, paragraph):
        self.paragraph = paragraph

    def fix_table_gridcol(self, table):
        cols = len(table.rows[0]._tr.tc_lst)
        for r in table.rows[1:]:
            row_cols = len(r._tr.tc_lst)
            if row_cols != cols:
                logger.warning(
                    'Table contains different cells count! Using bigger value')
            if row_cols > cols:
                cols = row_cols
        grid = CT_TblGrid()
        grid.tag = '{http://schemas.openxmlformats.org/wordprocessingml/' \
                   '2006/main}tblGrid'
        for i in range(cols):
            grid.add_gridCol()
        table._tbl.append(grid)

    def check_table_tags(self, table):
        try:
            table._tbl.tblGrid
        except InvalidXmlError:
            self.fix_table_gridcol(table)

    def del_table(self, table):
        el = table._element
        el.getparent().remove(el)
        table._tbl = table._element = None

    def process_table(self, table, table_data, is_symmetric, possible_header):
        # No data for postprocessing - usual processing
        if not is_symmetric:
            return self.usual_table_processing(table_data)
        # Two-column table postprocessing
        elif len(table_data[0]) == 2:
            return self.two_col_postprocessing(
                table_data,
                possible_header and self.has_horizontal_header(table, table_data)
            )
        # Case when there's empty first cell (table has header)
        elif len(table_data[0][0]) == 0:
            return self.header_type_1_postprocessing(table_data)
        # Case when there's horizontal header with different cells format
        elif possible_header and self.has_horizontal_header(table, table_data):
            return self.header_type_2_postprocessing(table_data)
        # Case when there's vertical header with different cells format
        elif possible_header and self.has_vertical_header(table):
            return self.header_type_2_postprocessing_vertical(table_data)
        # No cases - process as usual table
        return self.usual_table_processing(table_data)

    def process_prev_tables(self):
        tables = self.get_prev_tables()
        for table in tables:
            table_data, is_symmetric, possible_header = self.get_content_for_postprocessing(table)
            self.process_table(table, table_data, is_symmetric, possible_header)
            self.del_table(table)

    def get_prev_tables(self):
        tables = []
        prev = self.paragraph._p.getprevious()
        # While there are prev elements and prev paragraph isn't reached
        while prev is not None and not isinstance(prev, CT_P):
            # prepend table if this is table
            if isinstance(prev, CT_Tbl):
                tables.insert(0, Table(prev, prev.getparent()))
            prev = prev.getprevious()
        return tables

    def get_content_for_postprocessing(self, table):
        '''
        If table requires postprocessing - return text-only array of cells
        '''
        content = []
        self.check_table_tags(table)
        highlight_matrix = []
        for row in table.rows:
            r = []
            highlight_row = []
            previous_text = None

            for cell in row.cells:
                text = cell.text.strip().rstrip('!,-.:;')
                text = text.split('\n')
                for i in range(len(text)-1):
                    if text[i] and text[i][-1] not in punctuation:
                        text[i] += ';'
                text = ' '.join(text)

                # Ignore duplicate cell text
                if text == previous_text:
                    continue
                previous_text = text

                # Replace dots (except decimal dots) with ;
                mutable_text = list(text)
                for i, c in enumerate(text[:-1]):
                    if c == '.' and text[i+1] == ' ':
                        mutable_text[i] = ';'
                text = ''.join(mutable_text)

                r.append(text)

                # Get cell's highlight color if any
                style = self.cell_style(cell)['p']
                parsed = BeautifulSoup(style, 'xml')
                if parsed.find('w:highlight'):
                    highlight_row.append(parsed.find('w:highlight')['w:val'])
                elif u'\u2612' in cell.text:
                    highlight_row.append('\u2612')
                else:
                    highlight_row.append(None)

            content.append(r)
            highlight_matrix.append(highlight_row)

        highlighted_content = process_highlight(content, highlight_matrix)

        # Remove empty starting cells in each row
        if not any([row[0] for row in highlighted_content if row]):
            for i in range(len(highlighted_content)):
                if highlighted_content[i]:
                    highlighted_content[i].pop(0)

        is_symmetric = all([len(row) == len(content[0]) for row in highlighted_content])
        if highlighted_content == content:
            possible_header = True
        else:
            possible_header = False

        return highlighted_content, is_symmetric, possible_header

    def usual_table_processing(self, table_data):
        for row in table_data:
            paragraph = []
            cell_num = 0
            for text in row:
                if not text:
                    continue
                if cell_num == 0:
                    paragraph.append(text)
                    cell_num = 1
                elif cell_num == 1:
                    paragraph.append(' %s %s' % (get_separator(row[0]), text))
                    cell_num = 2
                else:
                    paragraph.append(' and %s' % text)

            self.insert_paragraph_before(''.join(paragraph))

    def insert_paragraph_before(self, text):
        new_par = self.paragraph.insert_paragraph_before(text)
        # assign styles
        new_par._p.style = self.paragraph._p.style

    def two_col_postprocessing(self, table_data, has_header):
        '''
        Postprocess two-column table

        If there are only two cols - write data in format 'ColA is/are ColB'
        '''

        if has_header:
            head1 = table_data[0][0]
            head2 = table_data[0][1]
            sep1 = get_separator(head1)
            sep2 = get_separator(head2)
            for row in table_data[1:]:
                self.insert_paragraph_before(
                    '%s %s %s and %s %s %s.' % (head1, sep1, row[0], head2, sep2, row[1])
                )
        else:
            for row in table_data:
                sep = get_separator(row[0])
                self.insert_paragraph_before(
                    '%s %s %s.' % (row[0], sep, row[1])
                )

    def header_type_1_postprocessing(self, table_data):
        '''
        Postprocess table with blank first cell (has header)

        Write data in format 'Col1 and HdrN is/are ColN'
        '''
        header = table_data[0]
        for row in table_data[1:]:
            for i, text in enumerate(row[1:]):
                sep = get_separator(header[i + 1])
                self.insert_paragraph_before(
                    '%s and %s %s %s.' % (row[0], header[i + 1], sep, text)
                )

    def cell_style(self, cell):
        # duplicate style
        style = docx.oxml.parse_xml(cell._tc.tcPr.xml)
        # remove meanless items (width/margin) and borders
        for i in style.getchildren():
            if 'tcW' in i.tag or 'tcMar' in i.tag or 'tcBorders' in i.tag:
                style.remove(i)

        p_style = BeautifulSoup(cell.paragraphs[0]._p.xml, 'xml')
        if p_style.t:
            # Remove text
            p_style.t.decompose()

        return {
            'cell': style.xml,
            'p': str(p_style)
        }

    def has_horizontal_header(self, table, table_data):
        '''
        Check if table has horizontal header
        '''

        # Table consists of only one row
        if len(table_data) == 1:
            return False

        header = table.rows[0]
        header_style = self.cell_style(header.cells[0])
        for cell in header.cells[1:]:
            if self.cell_style(cell) != header_style:
                return False
        for row in table.rows[1:3]:
            for cell in row.cells:
                if self.cell_style(cell) == header_style:
                    return False
        return True

    def header_type_2_postprocessing(self, table_data):
        '''
        Postprocess table horizontal header

        Write data in format 'Hdr1 is/are Col1 and ... and HdrN is/are ColN'
        '''
        header = table_data[0]

        separators = [get_separator(h) for h in header]

        for row in table_data[1:]:
            par_items = []
            for i, text in enumerate(row):
                par_items.append('%s %s %s' % (header[i], separators[i], text))

            text = '; and '.join(par_items)
            self.insert_paragraph_before(text + '.')

    def has_vertical_header(self, table):
        '''
        Check if table has vertical header
        '''
        header = table.rows[0]
        header_style = self.cell_style(table.rows[0].cells[0])
        for row in table.rows:
            if self.cell_style(row.cells[0]) != header_style:
                return False
            for cell in header.cells[1:]:
                if self.cell_style(cell) == header_style:
                    return False
        return True

    def header_type_2_postprocessing_vertical(self, table_data):
        '''
        Postprocess table vertical header

        Write data in format 'Hdr1 is/are Col1 and ... and HdrN is/are ColN'
        '''
        for i in range(1, len(table_data[0])):
            par_items = []
            for row in table_data:
                sep = get_separator(row[0])
                par_items.append('%s %s %s' % (row[0], sep, row[i]))
            text = ' and '.join(par_items)
            self.insert_paragraph_before(text + '.')


class lazy_property(object):
    def __init__(self, fget):
        self.fget = fget
        self.func_name = fget.__name__

    def __get__(self, obj, cls):
        if obj is None:
            return None
        value = self.fget(obj)
        setattr(obj, self.func_name, value)
        return value


class ParagraphDescription(object):
    TITLE_STYLES = ['heading', 'title', 'header']
    endstring_punc = (
        '!', '.', '?',
    )
    relstring_punc = (
        ':', ';'
    )
    spec_symbols = (
        '$', '%', '*', '@', '^', '+'
    )
    numbers = '0123456789'
    conjunctions = ('and', 'in', 'the', 'as', 'of', 'or')
    re_num_list = re.compile(u'^\(?(\d\.?)+\)?\s*[A-Za-z]+.*[^.]$')

    def __init__(self, paragraphs, index=None):
        # index can be '0'
        self.paragraph = None if index is None else paragraphs[index]
        self.paragraphs = paragraphs
        self.text = getattr(self.paragraph, 'text', None)
        self.index = index
        self.text_strip = self.text.strip() if self.text else ''

    @staticmethod
    def _rm_chars(text, to_remove):
        try:
            return text.translate({ord(c): None for c in to_remove})
        except TypeError:
            return text.translate(None, to_remove)

    def _remove_conjunctions(self, text):
        return ' '.join(t for t in text.split() if t not in self.conjunctions)

    def _remove_punctuation(self, text):
        return self._rm_chars(text, punctuation)

    def _remove_nums(self, text):
        return self._rm_chars(text, '1234567890.')

    def _split(self, text):
        ret = []
        for word in text.split():
            if word.startswith(('(', '\'', '"', '`')):
                ret += [word[0], word[1:]]
                continue
            elif word.endswith((')', '\'', '"', '`')):
                ret += [word[0:-1], word[-1]]
                continue
            ret.append(word)
        return ret

    def _remove_punc_nums(self):
        text = self._rm_chars(self.text, punctuation + self.numbers)
        return ' '.join(text.split())

    @lazy_property
    def length(self):
        return len(' '.join(self.text.split()))

    @lazy_property
    def uppercase(self):
        text = self._remove_punc_nums()
        if not len(text.split()):
            return False
        return all(t.isupper() for t in self._split(text))

    @lazy_property
    def camelcase(self):
        if self.word_count == 1:
            return False
        text = self._remove_punc_nums()
        text = self._remove_conjunctions(text)
        words = self._split(text)
        if len(words) == 1:
            return False
        title_count = sum(1 for word in words if word.istitle() or word.isupper())
        # ratio between uppers/titles count and total words count more then 60%
        if title_count == 0:
            return False
        return (float(title_count) / len(words) * 100) > 65.0

    @lazy_property
    def start_lower(self):
        if self.listing:
            text = RE_LIST.sub('', self.text)
        else:
            text = self.text
        return text.strip().startswith(tuple(ascii_lowercase))

    @lazy_property
    def last_symbol(self):
        if not self.text:
            return ''
        return self.text[-1]

    @lazy_property
    def listing(self):
        if not self.text:
            return False
        return self.re_num_list.match(self.text_strip) or self.paragraph._p.pPr and self.paragraph._p.pPr.numPr

    @lazy_property
    def next_non_empty(self):
        index = self.index
        while True:
            index += 1
            try:
                next_par = self.paragraphs[index]
                if bool(next_par.text.strip()):
                    return ParagraphDescription(self.paragraphs, index)
            except IndexError:
                return ParagraphDescription(self.paragraphs)
            except AttributeError:
                continue

    @lazy_property
    def prev_non_empty(self):
        index = self.index
        while True:
            index -= 1
            try:
                prev = self.paragraphs[index]
                if prev.text:
                    return ParagraphDescription(self.paragraphs, index)
            except IndexError:
                return ParagraphDescription(self.paragraphs)
            except AttributeError:
                continue

    @lazy_property
    def word_count(self):
        return len(re.findall(u'[A-Za-z]+', self.text))

    @property
    def next_paragraph(self):
        index = self.index + 1
        try:
            self.paragraphs[index]
        except IndexError:
            next_par = None
        else:
            next_par = ParagraphDescription(self.paragraphs, index)
        return next_par

    @property
    def prev_paragraph(self):
        index = self.index - 1
        try:
            self.paragraphs[index]
        except IndexError:
            prev_par = None
        else:
            prev_par = ParagraphDescription(self.paragraphs, index)
        return prev_par

    @lazy_property
    def lead_spaces(self):
        return len(re.search(u'^\s*', self.text).group(0))

    @lazy_property
    def number_count(self):
        return len(re.findall(u'[0-9]+', self.text))

    def has_any_style(self, styles):
        return any(s in self.paragraph.style.name.lower() for s in styles)

    def centered(self):
        return self.paragraph.alignment in [1, 2]

    @lazy_property
    def listing_position(self):
        is_next_listing = getattr(self.next_non_empty, 'listing', None)
        is_prev_listing = getattr(self.prev_non_empty, 'listing', None)
        if self.listing:
            if is_next_listing and not is_prev_listing:
                return 'top'
            elif is_prev_listing and is_next_listing:
                return 'inside'
            elif not is_next_listing and is_prev_listing:
                return 'last'

    @lazy_property
    def similar_outside(self):
        try:
            prev_text = self.prev_non_empty.text or ''
        except AttributeError:
            prev_text = ''
        try:
            next_text = self.next_non_empty.text or ''
        except AttributeError:
            next_text = ''
        curr_words = self.text.split()
        if prev_text.startswith(curr_words[0]) or next_text.startswith(curr_words[0]):
            return True

    def list_item_before_text(self):
        if not self.listing or self.next_non_empty.startswith_lower:
            return False
        return not self.next_non_empty.listing

    def lead_upper_then_camel(self):
        """
        For case like "SCHEDULE 1 - Proposed Worker Definitions"
        """
        if self.word_count == 1:
            return False
        text = self._remove_punc_nums()
        return bool(re.match(u'^[A-Z]+\s(?:[A-Z][a-z]+\s?)+$', text))

    @lazy_property
    def is_page(self):
        patterns = (
            re.compile(u'[Pp]age[:]?\s*\d+$'),
            re.compile(u'\d+\s?of\s?\d+'),
            re.compile(u'^\s*\d+\s*$')
        )
        for pattern in patterns:
            if bool(pattern.search(self.text.rstrip())):
                return True

    @lazy_property
    def uppers_count(self):
        return len(re.findall(u'[A-Z]{2,}', self.text))

    @lazy_property
    def lowers_count(self):
        return len(re.findall(u'([A-Z][a-z]|[a-z]){2,}', self.text))

    @lazy_property
    def white_spaces_count(self):
        return len(re.findall(u'(\s{4,})', self.text))

    @property
    def has_lead_whitespace(self):
        return self.text_strip and bool(re.match(u'^\s+', self.text))

    @lazy_property
    def startswith_upper(self):
        return self.text_strip.startswith(tuple(ascii_uppercase))

    @lazy_property
    def startswith_lower(self):
        return self.text_strip.startswith(tuple(ascii_lowercase))

    def list_elem_related_to_text_below(self):
        if any((not self.listing, self.next_non_empty.text is None, self.next_non_empty.startswith_upper)):
            return False
        elif self.text.rstrip().endswith(self.conjunctions) or self.next_non_empty.startswith_lower:
            return True
        next_match = re.match(u'^\s+', self.next_non_empty.text)
        cur_match = self.re_num_list.match(self.text)
        if not next_match or self.top_level_num_list_elem() or not cur_match:
            return False
        num_white_count = len(cur_match.group(0))
        num_white_count_next = len(next_match.group(0))
        return num_white_count <= num_white_count_next

    def top_level_num_list_elem(self):
        if not self.listing or not self.next_non_empty.listing:
            return False
        elif bool(re.match(u'\s*[\d\.]+\.0.*$', self.text)):
            return True
        re_extract_num = re.compile(u'^(\d+\.?)+')
        try:
            cur_num = re_extract_num.match(self.text_strip).group(0)
            next_num = re_extract_num.match(self.next_non_empty.text_strip).group(0)
        except AttributeError:
            return False
        else:
            return len(cur_num) < len(next_num)

    def is_include_spec_symbols(self):
        return any(s in self.text for s in self.spec_symbols)

    def camel_rule(self, listing=False):
        if self.text.endswith('.') or self.white_spaces_count >= 2:
            return False
        if listing:
            return self.camelcase and self.listing
        return self.camelcase and not self.listing

    def os_path(self):
        return bool(re.match(r'^[A-Z]:(\\.*){2,}', self.text_strip))

    def single_symbol(self):
        return self.text_strip in list(punctuation)

    def prev_empty(self):
        is_empty = any((
            not self.prev_paragraph, not self.prev_paragraph.text, not self.prev_paragraph.text_strip
        ))
        return is_empty

    def is_toc(self, probability='high'):
        patterns = []
        if probability == 'high':
            patterns.extend([TOCCleanup.TOC_ITEM_1, TOCCleanup.TOC_ITEM_4])
        else:
            patterns.extend([TOCCleanup.TOC_ITEM_3])
        keywords = TOCCleanup.TOC_NAMES
        for pattern in patterns:
            if pattern.match(self.text_strip):
                return True
        return any(k in self.text.lower() for k in keywords)

    def hft_points(self):
        return [
            (self.centered, 1.0),
            (self.os_path, 1.0),
            (self.single_symbol, 1.0),
            (lambda: self.length > 120, -0.9),
            (lambda: self.text.count('.') > 4, -1.0),
            (self.is_toc, -0.9),
            (lambda: self.is_toc(probability='low'), -0.4),
            (lambda: self.has_any_style(self.TITLE_STYLES), 0.9),
            (lambda: not self.uppercase and self.listing_position in ['inside', 'last'], -0.9),
            (lambda: 120 > self.length > 100, -0.7),
            (lambda: 100 > self.length > 80, -0.5),
            (lambda: self.text.endswith(self.conjunctions), -0.9),
            (self.prev_empty, 0.8),
            (lambda: self.uppercase, 2.0),
            (lambda: self.is_page, 2.0),
            (lambda: self.start_lower and self.lead_spaces <= 2, -0.1),
            (lambda: self.listing and self.start_lower, -0.7),
            (lambda: self.text.endswith(self.endstring_punc), -0.1),
            (lambda: self.text.endswith(self.relstring_punc), -1.3),
            (lambda: any(s in self.text_strip[:-1] for s in self.relstring_punc), -0.4),
            (lambda: self.white_spaces_count > 2, -1.3),
            (self.list_elem_related_to_text_below, -0.7),
            (self.is_include_spec_symbols, -0.5),
            (self.top_level_num_list_elem, 1.7),
            (lambda: 80 > self.length > 60, 0.1),
            (lambda: self.uppers_count > self.lowers_count, 1.0),
            (self.lead_upper_then_camel, 0.9),
            (lambda: self.word_count == 1 and self.text.startswith(ascii_uppercase), 1.0),
            (lambda: self.camel_rule(listing=False), 0.8),
            (lambda: self.camel_rule(listing=True), 1.0),
            (self.list_item_before_text, 0.8),
            (lambda: self.lead_spaces >= self.length, 0.7),
            (lambda: 61 > self.length, 0.2),
            (lambda: self.listing and self.word_count <= 2, 0.8),
        ]


class TitleCleanup:
    def __init__(self, docxpath):
        self.doc = docx.Document(docxpath) if isinstance(docxpath, basestring) else docxpath
        self.paragraphs = filter(lambda par: bool(par.text.strip()), self.doc.paragraphs)
        self.docxpath = docxpath

    def calculate(self, par_desc):
        total_points = 0
        for index, (func, point) in enumerate(par_desc.hft_points()):
            if func():
                total_points += point
            if total_points >= 1.0:
                return True
            elif total_points <= -1.0:
                return False

    def run(self):
        for paragraph_index, paragraph in [p for p in enumerate(self.paragraphs) if p[1].text]:
            par_desc = ParagraphDescription(self.paragraphs, paragraph_index)
            if self.calculate(par_desc):
                del_paragraph(paragraph)
        return self.doc


class HFCleanup:
    """
    Removes docx header & footer
    """
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        self.zipfile = zipfile.ZipFile(docx_path)
        self.extract_path = tempfile.mkdtemp()
        self.zipfile.extractall(self.extract_path)
        self.zipfile.close()

    def get_file(self, xmlpath):
        file_path = os.path.join(self.extract_path, xmlpath)
        return open(file_path).read()

    def get_xml_tree(self, xmlpath):
        return fromstring(self.get_file(xmlpath))

    def find_in(self, root, nodes, ns):
        cur = root
        for node in nodes[0:-1]:
            cur = cur.find(node, ns)
        return cur.findall(nodes[-1], ns)

    def clean_document_xml(self):
        zip_path = 'word/document.xml'
        tree = self.get_xml_tree(zip_path)
        for header in tree.findall('.//w:headerReference', self.ns):
            header.getparent().remove(header)
        for footer in tree.findall('.//w:footerReference', self.ns):
            footer.getparent().remove(footer)
        self._write(tree, zip_path)

    def cleanup_temp(self):
        try:
            shutil.rmtree(self.extract_path)
        except OSError:
            pass

    def _write(self, tree, xmlpath):
        with open(os.path.join(self.extract_path, xmlpath), 'w') as f:
            f.write(tostring(tree))

    def _make_docx(self):
        out_path = os.path.join(self.extract_path, os.path.basename(self.docx_path))
        docx = zipfile.ZipFile(out_path, 'w')
        filenames = self.zipfile.namelist()
        for filename in filenames:
            docx.write(os.path.join(self.extract_path, filename), filename)
        return out_path

    def run(self):
        self.clean_document_xml()
        filepath = self._make_docx()
        doc = docx.Document(filepath)
        self.cleanup_temp()
        return doc


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
