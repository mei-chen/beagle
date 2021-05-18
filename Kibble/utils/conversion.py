"""
    conversion
    ==========

    A library to convert documents to docx standard.
"""

import logging
import os
import re
import string
import subprocess
import sys
from io import BytesIO
from subprocess import call
from zipfile import ZipFile, ZIP_DEFLATED

import chardet
from bs4 import BeautifulSoup
from bs4.element import NavigableString
from django.conf import settings
from docx import Document
from pdf2docx import Converter
import ocrmypdf


logger = logging.getLogger(__name__)

page_rgx = re.compile(r'^[Pp]age \d+\s*$')


class DocumentSizeOverLimitException(Exception):
    pass


# terminate cmd after 60 secs, kill after 60+30 secs
CALL_TIMEOUT = ['timeout', '-k', '30', '60']


def strings(f, min=4):
    """
    Helper function to iterate through a file binary and pull out strings
    called by requires_ocr()
    """
    result = ""
    for c in f.read():
        c = str(c)
        if c in string.printable:
            result += c
            continue
        if len(result) >= min:
            yield result
        result = ""


def requires_ocr(f):
    """
    Helper function to look for 'FontName' in a binary file
    indicating text presence and hence no necessity for OCR
    """
    is_ocr = True
    for s in strings(f):
        if 'FontName' in s:
            is_ocr = False
            break
    return is_ocr


def pdf_to_docx(filename, ocr):
    """
    Transcodes a .pdf file to docx via the EasyPdfCloud
    :param filename: path to locate the .pdf file
    :returns: path to the docx version of the file
    """

    upload = os.path.abspath(filename)
    if not os.path.isfile(upload):
        raise Exception('File %s does not exist.' % upload)
    pdf_convert(upload, ocr)
    docxpath = get_filename_wo_ext(filename) + '.docx'
    return docxpath


def doc_to_docx(filename):
    """
    Transcodes a .pdf file to docx via the EasyPdfCloud
    :param filename: path to locate the .pdf file
    :returns: path to the docx version of the file
    """
    upload = os.path.abspath(filename)
    doc_convert(upload)

    docxpath = get_filename_wo_ext(filename) + '.docx'
    return docxpath


def contentsdecode(text):
    found_encoding = chardet.detect(text)['encoding']
    dec_contents = text.decode(found_encoding or 'utf-8', errors='ignore')
    return dec_contents


def filedecode(filename):
    """
    Determines encoding of a file, and outputs utf-8 encoded plaintext
    contents
    :param filename: path to locate the file
    :returns: plaintext contents of the file
    """
    with open(filename, 'rb') as f:
        file_contents = f.read()
        return contentsdecode(file_contents)


def txt_to_docx(filename):
    """
    Transcodes a .txt file to docx using python-docx
    :param filename: path to locate the .txt file
    :returns: path to the docx version of the file
    """
    content = filedecode(filename)
    docxpath = get_filename_wo_ext(filename) + '.docx'
    document = Document()
    for line in content.replace(u'\xa0', ' ').splitlines():
        paragraph = document.add_paragraph(line)
        paragraph.style = document.styles['Normal']
    document.save(docxpath)
    return docxpath


def docx_to_txt(filename):
    """
    Transcodes a .doc file to txt using docx2txt
    :param filename: path to locate the .docx file
    :returns: path to the txt version of the file
    """
    txtname = get_filename_wo_ext(filename) + '.txt'
    docx2txt = os.path.join(
        settings.BASE_DIR, 'utils', 'docx2txt', 'docx2txt.pl')
    if call(CALL_TIMEOUT + [docx2txt, filename, txtname]) != 0:
        return None
    return txtname


def html_to_docx(filename):
    pdfpath = get_filename_wo_ext(filename) + '.pdf'
    html_to_pdf(filename, pdfpath)
    ocr_needed = requires_ocr(open(pdfpath, 'rb'))
    docx = pdf_to_docx(pdfpath, ocr_needed)

    return docx


def xml_to_docx(filename):
    docxpath = get_filename_wo_ext(filename) + '.docx'
    document = Document()
    parsed = BeautifulSoup(contentsdecode(open(filename).read()), 'xml')
    lines = [child.strip() for child in parsed.recursiveChildGenerator()
             if isinstance(child, NavigableString) if child.strip()]
    text = '\n\n'.join(lines)

    html_parsed = BeautifulSoup(text.encode('utf-8'), 'html.parser')

    for line in html_parsed.text.replace(u'\xa0', ' ').splitlines():
        if page_rgx.match(line):
            continue
        paragraph = document.add_paragraph(line)
        paragraph.style = document.styles['Normal']
    document.save(docxpath)

    return docxpath


def docx_libreoffice(filename):
    """
    Transcodes file to docx using libreoffice
    :param filename: path to locate source file
    :returns: path to the docx version of the file
    """
    docx_name = get_filename_wo_ext(os.path.basename(filename)) + '.docx'
    dest_dir = '/tmp'
    if call(CALL_TIMEOUT + [
            'libreoffice', '--headless', '--convert-to', 'docx',
            '--outdir', dest_dir, filename]) != 0:
        return None

    docx_name = os.path.join(dest_dir, docx_name)
    if os.path.isfile(docx_name):
        return docx_name


def compress_to_zip(files):
    """
    Creates filelike object with ZIP-compressed files
    :param files: list of files to compress
    :returns: filelike object with ZIP data
    """
    content = BytesIO()
    with ZipFile(content, mode='w', compression=ZIP_DEFLATED) as zf:
        for f in files:
            if not f:
                continue
            name = os.path.basename(f.name)
            if isinstance(name, bytes):
                name = name.decode('utf-8')
            try:
                f.seek(0)
                zf.writestr(name, f.read().decode('utf-8'))
                f.seek(0)
            except:
                logger.warn(
                    "Can't add {} to archive".format(name), exc_info=True)
    return content


class InvalidDocumentTypeException(Exception):
    pass


def get_filename_wo_ext(filepath):
    name, ext = os.path.splitext(filepath)
    return name


def docx_to_pdf(folder, source):
    args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', folder, source]
    pp = subprocess.Popen(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    out, err = pp.communicate()

    if pp.returncode:
        logger.error(err.strip())
    else:
        logger.info(out.strip())


def html_to_pdf(html_path, pdf_path):
    args = ['wkhtmltopdf', '--encoding', 'utf-8', html_path, pdf_path]
    pp = subprocess.Popen(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    out, err = pp.communicate()

    if pp.returncode:
        logger.error(err.strip())
    else:
        logger.info(out.strip())


def pdf_convert(input_filename, ocr):
	"""
	Given in file name, check if ocr is needed
	If so, used ocrmypdf to add text layer
	then convert to docx
	"""

	#if ocr:
	pdf_ocr(input_filename)
	
	output_filename = '.'.join(os.path.abspath(input_filename).split('.')[:-1]) + '.docx'

	# convert pdf to docx
	cv = Converter(input_filename)
	cv.convert(output_filename, start=0, end=None)
	cv.close()

	return output_filename


def pdf_ocr(input_filename):

    args = ["/".join(sys.executable.split("/")[:-1]) + "/ocrmypdf", '--skip-text', '--output-type', 'pdf', input_filename, input_filename]
    pp = subprocess.Popen(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    out, err = pp.communicate()

    if pp.returncode:
        logger.error(err.strip())
    else:
        logger.info(out.strip())


def doc_convert(input_filename):
    output_directory = os.path.dirname(input_filename)
    subprocess.call(['soffice', '--headless', '--convert-to', 'docx', '--out-dir', output_directory, input_filename])