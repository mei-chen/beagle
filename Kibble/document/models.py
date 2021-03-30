from __future__ import unicode_literals

import csv
import json
import logging
import os
import tempfile
import uuid
from io import BytesIO, StringIO

import docx
import numpy as np
from django.conf import settings
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.db import models
from django.db.models.fields.files import FieldFile
from django.db.models.signals import pre_delete
from django.dispatch import receiver
from jsonfield import JSONField

from realtime.notify import NotificationManager
from utils.cleanup import (
    is_not_newline, del_paragraph,
    TitleCleanup, TOCCleanup, TablePostprocessing, ListCleanup, HFCleanup
)
from utils.conversion import docx_to_txt, docx_libreoffice
from utils.sentence_vector.api import SentenceVectorAPI

logger = logging.getLogger(__name__)


def make_upload_path(instance, fname):
    if instance.batch is not None:
        name = instance.batch.name.replace(
            ' ', '_')[:settings.BATCH_NAME_TRUNCATE_CHARS]
        predicate = '{}_{}'.format(instance.batch.id, name)
    else:
        predicate = 'initially_unassigned'
    path = os.path.join('batch_docs', predicate, fname)
    return path


class Document(models.Model):

    class Meta:
        ordering = ['id']

    name = models.CharField(max_length=300)
    uuid = models.UUIDField(default=uuid.uuid4)
    created_at = models.DateTimeField(auto_now_add=True)
    content_file = models.FileField(
        null=True, max_length=400, upload_to=make_upload_path)
    text_file = models.FileField(
        null=True, max_length=400, upload_to=make_upload_path)
    source_file = models.ForeignKey(
        'portal.File', null=True, related_name='documents',on_delete=models.CASCADE)
    origin = models.OneToOneField(
        'self', on_delete=models.CASCADE, null=True, related_name='cleaned')

    def __str__(self):
        return self.name

    def get_doc(self):
        return docx.Document(self.content_file)

    def get_text_lines(self):
        if not self.text_file:
            return []
        self.text_file.seek(0)
        lines = self.text_file.readlines()
        self.text_file.seek(0)
        return [l.decode('utf-8') for l in lines]

    def store_doc(self, doc):
        io = BytesIO()
        doc.save(io)
        io.seek(0)
        self.store_content(io.read())

    def store_content(self, content):
        cf = ContentFile(content)
        cf.name = self.content_file.name
        self.content_file = cf
        self.text_file = None
        self.save()

    @classmethod
    def get_file(cls, field):
        f = tempfile.NamedTemporaryFile(delete=True)
        if not field:
            return f
        field.seek(0)
        f.write(field.read())
        field.seek(0)
        f.flush()
        return f

    def get_content_file(self):
        return self.get_file(self.content_file)

    def get_text_file(self):
        return self.get_file(self.text_file)

    def save(self, *args, **kwargs):
        needs_conversion = bool(self.content_file)
        if self.pk:
            old_data = Document.objects.get(pk=self.pk)
            needs_conversion = needs_conversion and (
                not self.text_file or
                old_data.content_file != self.content_file
            )
            if needs_conversion:
                # cleanup - delete old files
                old_data.content_file.delete(False)
                old_data.text_file.delete(False)
        else:
            needs_conversion = needs_conversion and \
                not self.text_file

        ret = super(Document, self).save(*args, **kwargs)
        if not needs_conversion:
            return ret

        f = self.get_content_file()
        try:
            abspath = docx_to_txt(f.name)
        finally:
            f.close()
        if abspath:
            with open(abspath, 'rb') as f:
                content = f.read()
            os.unlink(abspath)
            cf = ContentFile(content)
            cf.name = os.path.splitext(
                os.path.basename(self.content_file.name))[0] + '.txt'
            self.text_file = cf

            # instance already created, we need to update it instead
            kwargs['force_insert'] = False
            return super(Document, self).save(*args, **kwargs)
        else:
            logger.warning(
                "Can't convert file %s to text" % self.content_file.name
            )
        return ret

    def make_copy(self):
        self.refresh_from_db()
        try:
            self.cleaned.delete()
        except self.DoesNotExist:
            pass

        text_file = None
        if self.text_file:
            text_file = ContentFile(self.text_file.read())
            text_file.name = 'cleaned_' + os.path.basename(self.text_file.name)

        content_file = None
        if self.content_file:
            content_file = ContentFile(self.content_file.read())
            content_file.name = 'cleaned_' + os.path.basename(
                self.content_file.name)
        copied_name = "Cleaned copy {}".format(self.name)
        # Temporary solution for unittests - righ truncate name if greater than 300 characters
        if len(copied_name) > 300:
            copied_name = copied_name[:300]
        copy = Document.objects.create(
            name=copied_name,
            text_file=text_file,
            content_file=content_file,
            origin=self
        )
        return copy

    def converyor_error(self, message, session):
        logger.warning(message)
        if session:
            NotificationManager.popup_notification(session, message, 'warning')

    def conveyor(self, tools, session=None):
        doc = self.make_copy()
        for index, tool in enumerate(tools):
            cleanup_tool = self.get_tool(tool=tool)
            if cleanup_tool is None:
                # Bad tool
                doc.converyor_error('Skip unknown tool "%s"' % tool, session)
            elif cleanup_tool(doc):
                # Tool succeed
                DocumentTag.objects.create(
                    name=tool, document=self, order=index)
            else:
                # Tool failed
                doc.converyor_error(
                    'Tool "%s" failed on %s!' % (tool, self.name), session)
        return doc

    def linebreakers_cleanup(self):
        doc = self.get_doc()
        prev_paragraph = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.rstrip(' \t')
            # save strip spaces at end
            paragraph.text = text
            # skip lists
            if ListCleanup.is_list_elem(paragraph):
                prev_paragraph = paragraph
                continue
            # strip spaces at start
            text = text.strip(' \t')
            # save stripped text
            paragraph.text = text
            # remove empty line
            if not text:
                del_paragraph(paragraph)
                continue
            if not prev_paragraph:
                prev_paragraph = paragraph
                continue
            # skip if this is newline
            if not is_not_newline(text, prev_paragraph.text):
                prev_paragraph = paragraph
                continue
            prev_text = ' '.join((prev_paragraph.text.rstrip(' \t'), text))
            prev_paragraph.text = prev_text
            del_paragraph(paragraph)
        self.store_doc(doc)
        return self

    def title_header_footer_cleanup(self):
        file = self.get_content_file()
        try:
            doc = HFCleanup(file.name).run()
        finally:
            file.close()
        doc = TitleCleanup(doc).run()
        self.store_doc(doc)
        return self

    def table_of_contents_cleanup(self):
        try:
            doc = TOCCleanup(self.content_file.name, self.get_doc()).run()
        except Exception as e:
            logger.warning(e)
            return None
        else:
            self.store_doc(doc)
        return self

    def bullet_points_and_listing_cleanup(self):
        tool = ListCleanup(self.get_doc(), self.get_text_lines())
        doc = tool.list_cleanup(False)
        self.store_doc(doc)
        return self

    def listing_rewriting_cleanup(self):
        tool = ListCleanup(self.get_doc(), self.get_text_lines())
        doc = tool.list_cleanup(True)
        self.store_doc(doc)
        return self

    def tables_cleanup(self):
        doc = self.get_doc()
        for paragraph in doc.paragraphs:
            postprocessor = TablePostprocessing(paragraph)
            postprocessor.process_prev_tables()
        self.store_doc(doc)
        return self

    def normalize_LO(self):
        file = self.get_content_file()
        try:
            new = docx_libreoffice(file.name)
        finally:
            file.close()
        if not new:
            logger.warning(
                "%s is NOT cleaned by libreoffice!" % self.content_file.name)
            return None
        with open(new, 'rb') as f:
            self.store_content(f.read())
        return self

    @staticmethod
    def get_tool(tool=None):
        tools = {
            'normalize (libreoffice)': Document.normalize_LO,
            'linebreakers': Document.linebreakers_cleanup,
            'title/header/footer': Document.title_header_footer_cleanup,
            'table of contents': Document.table_of_contents_cleanup,
            'bullet points and listing':
                Document.bullet_points_and_listing_cleanup,
            'listing rewriting': Document.listing_rewriting_cleanup,
            'tables': Document.tables_cleanup,
        }
        return tools.get(tool) if tool is not None else tools

    def cleanup(self, tools, session=None):
        if tools:
            self.purge_tags()
            self.conveyor(tools, session)
        return self

    @property
    def has_sentences(self):
        return self.sentences.exists()

    def get_csv(self):
        if not self.has_sentences:
            return None
        filelike = StringIO()
        writer = csv.writer(filelike, quoting=csv.QUOTE_NONNUMERIC)
        for sentence in self.sentences.all():
            writer.writerow([sentence.text])
        return filelike.getvalue()

    def get_json(self):
        return json.dumps([
            sentence.text for sentence in self.sentences.all()
        ])

    @property
    def cleaned_txt_url(self):
        if not hasattr(self, 'cleaned'):
            return None
        return self.cleaned.text_file.url

    @property
    def cleaned_docx_url(self):
        if not hasattr(self, 'cleaned'):
            return None
        return self.cleaned.content_file.url

    def purge_tags(self):
        DocumentTag.objects.filter(document=self).delete()

    @property
    def batch(self):
        inst = self.origin or self
        return inst.source_file and inst.source_file.batch


class DocumentTag(models.Model):
    document = models.ForeignKey(Document, related_name='tags',on_delete=models.CASCADE)
    name = models.CharField(max_length=300)
    order = models.IntegerField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.name


class Sentence(models.Model):
    uuid = models.UUIDField(default=uuid.uuid4)
    created_at = models.DateTimeField(auto_now_add=True)
    document = models.ForeignKey(Document, related_name='sentences',on_delete=models.CASCADE)
    text = models.TextField()

    # The sentence in a vectorized form. The form can be absent or even change
    # with time, so use the `vector` property instead in order to ensure both
    # lazy initialization and proper loading into a NumPy/SciPy vector.
    vectorization = JSONField(null=True, blank=True)

    class Meta:
        ordering = ['id']

    @property
    def vector(self):
        # Lazy vectorization    
        if self.vectorization is None:
            vector = SentenceVectorAPI(self.text).vectorize()
            if vector is not None:
                # `numpy.ndarray` "is not JSON serializable"
                self.vectorization = list(vector)
                self.save()
        return (
            None if self.vectorization is None else np.array(self.vectorization)
        )

    @vector.setter
    def vector(self, sentence_vector):
        # `numpy.ndarray` "is not JSON serializable"
        self.vectorization = None if sentence_vector is None else list(sentence_vector)
        self.save()

    def __str__(self):
        return self.text


class PersonalData(models.Model):
    uuid = models.UUIDField(default=uuid.uuid4)
    created_at = models.DateTimeField(auto_now_add=True)
    type = models.CharField(max_length=255)
    text = models.CharField(max_length=255)
    location = models.CharField(max_length=255)
    document = models.ForeignKey(Document,on_delete=models.CASCADE)
    user = models.ForeignKey(User, null=True,on_delete=models.CASCADE)
    selected = models.BooleanField(default=True)

    def __str__(self):
        return self.text

    def save(self, *args, **kwargs):
        # Discard too short entries
        if len(self.text) < 2:
            return

        self.user = self.document.source_file.batch.owner
        if not self.pk:
            if self.type not in self.user.profile.personal_data_types:
                self.user.profile.personal_data_types[self.type] = (True, 'string')
                self.user.profile.save()

            self.selected = self.user.profile.personal_data_types[self.type][0]

        super(PersonalData, self).save(*args, **kwargs)

    @property
    def batch(self):
        if self.document.source_file:
            return self.document.source_file.batch


class CustomPersonalData(models.Model):
    uuid = models.UUIDField(default=uuid.uuid4)
    created_at = models.DateTimeField(auto_now_add=True)
    type = models.CharField(max_length=255)
    text = models.CharField(max_length=255)
    is_rgx = models.BooleanField(default=False)
    user = models.ForeignKey(User, null=True,on_delete=models.DO_NOTHING,)
    selected = models.BooleanField(default=True)

    def __str__(self):
        return self.text

    def save(self, *args, **kwargs):
        if self.type not in self.user.profile.personal_data_types:
            self.user.profile.personal_data_types[self.type] = (True, 'string')
            self.user.profile.save()
        self.selected = self.user.profile.personal_data_types[self.type][0]

        super(CustomPersonalData, self).save(*args, **kwargs)


def file_cleanup(sender, instance, *args, **kwargs):
    '''
    Deletes the file(s) associated with a model instance.
    '''
    for field_name, _ in instance.__dict__.items():
        field = getattr(instance, field_name)
        if issubclass(field.__class__, FieldFile) and field.name:
            try:
                field.delete(save=False)
            except:
                logger.warning(
                    'Delete: omitting file {}'.format(field.name)
                )


pre_delete.connect(file_cleanup)


@receiver(models.signals.post_save, sender=Document)
def gather_personal_data(sender, instance, *args, **kwargs):
    """
    Gathers all found personal data from a converted document and
    creates DocumentPersonalData objects for each entry.
    """
    from utils.personal_data import find_personal_data

    if (not instance.content_file or
        not instance.source_file or
        not instance.source_file.batch or
        not instance.source_file.batch.owner or
        not instance.source_file.batch.owner.profile or
        not instance.source_file.batch.owner.profile.auto_gather_personal_data or
        PersonalData.objects.filter(document=instance).exists()):
            return

    pd = find_personal_data(instance)
    for pd_type, text, location in pd:
        PersonalData.objects.create(
            document=instance,
            type=pd_type,
            text=text,
            location=location
        )

    batch = instance.source_file.batch
    if not batch.personal_data_gathered:
        batch.personal_data_gathered = True
        batch.save()
